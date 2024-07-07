from docx.api import Document
import os, shutil
import spacy
import numpy as np
from rake_nltk import Rake
import nltk
from nltk.corpus import stopwords
from nltk import punkt
import string 
import sys

nlp = spacy.load("en_core_web_lg")

path = sys.argv[1]

mkey = ""
pkey = ""
ckey = ""
file = Document("/Users/tanvi/Documents/Project/Dataset/MathKeywords2.docx")
for p in file.paragraphs:
    mkey += p.text +" "

file = Document("/Users/tanvi/Documents/Project/Dataset/chemKeywords.docx")
for p in file.paragraphs:
    ckey += p.text+ " "

file = Document("/Users/tanvi/Documents/Project/Dataset/PhysKeywords.docx")
for p in file.paragraphs:
    pkey += p.text + " "

mkey = mkey.lower()
pkey = pkey.lower()
ckey = ckey.lower()


mkey = nlp(mkey)
pkey = nlp(pkey)
ckey = nlp(ckey)

test = Document(path)
r = Rake()
content = ""
for p in test.paragraphs:
    content+= p.text
    content = content.replace("\n"," ")
con = ""
for i in content:
    if(ord(i) == 32 or (ord(i) > 64 and ord(i) < 91) or (ord(i) > 96 and ord(i) < 123)):
        con+=i
    else:
        con+=" "

con = nlp(con)
con_noun = " ".join([token.lemma_ for token in con if token.pos_ == "NOUN"])

conkey = nlp(con_noun)

c = conkey.similarity(ckey)
m = conkey.similarity(mkey)
p = conkey.similarity(pkey)

if m>p and m>c:
	shutil.move(path, "/Users/tanvi/Documents/Project/Math")
elif p>m and p>c:
	shutil.move(path, "/Users/tanvi/Documents/Project/Physics")
else:
	shutil.move(path, "/Users/tanvi/Documents/Project/Chemistry")

print("Done")
